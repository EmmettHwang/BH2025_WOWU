"""
문서 로더 모듈
PDF, DOCX, TXT 파일을 로드하고 텍스트를 추출합니다.
"""

import os
from typing import List, Dict
from pathlib import Path
import PyPDF2
import docx

# LangChain imports - 버전 호환성 처리
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document


class DocumentLoader:
    """문서 로드 및 청킹 클래스"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Args:
            chunk_size: 청크 크기 (토큰 수)
            chunk_overlap: 청크 간 겹침 (토큰 수)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> str:
        """PDF 파일에서 텍스트 추출"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            print(f"❌ PDF 로드 실패: {file_path}, 오류: {e}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트 추출"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            print(f"❌ DOCX 로드 실패: {file_path}, 오류: {e}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """TXT 파일에서 텍스트 추출"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            print(f"❌ TXT 로드 실패: {file_path}, 오류: {e}")
            return ""
    
    def load_document(self, file_path: str, metadata: Dict = None) -> List[Document]:
        """
        파일 확장자에 따라 적절한 로더로 문서 로드 후 청킹
        
        Args:
            file_path: 파일 경로
            metadata: 메타데이터 (예: {"source": "강의록", "date": "2024-01-01"})
            
        Returns:
            청크로 나뉜 Document 리스트
        """
        file_ext = Path(file_path).suffix.lower()
        
        # 파일 타입에 따라 텍스트 추출
        if file_ext == '.pdf':
            text = self.load_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.load_docx(file_path)
        elif file_ext == '.txt':
            text = self.load_txt(file_path)
        else:
            print(f"⚠️ 지원하지 않는 파일 형식: {file_ext}")
            return []
        
        if not text:
            print(f"⚠️ 빈 문서: {file_path}")
            return []
        
        # 메타데이터 설정
        if metadata is None:
            metadata = {}
        
        metadata.update({
            'source': os.path.basename(file_path),
            'file_path': file_path,
            'file_type': file_ext
        })
        
        # 텍스트를 청크로 분할
        chunks = self.text_splitter.split_text(text)
        
        # Document 객체로 변환
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = metadata.copy()
            chunk_metadata['chunk_id'] = i
            chunk_metadata['total_chunks'] = len(chunks)
            
            doc = Document(
                page_content=chunk,
                metadata=chunk_metadata
            )
            documents.append(doc)
        
        print(f"✅ 문서 로드 완료: {os.path.basename(file_path)} ({len(documents)}개 청크)")
        return documents
    
    def load_directory(self, directory_path: str, metadata: Dict = None) -> List[Document]:
        """
        디렉토리 내 모든 문서 로드
        
        Args:
            directory_path: 디렉토리 경로
            metadata: 공통 메타데이터
            
        Returns:
            모든 문서의 청크 리스트
        """
        all_documents = []
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                file_ext = Path(file_path).suffix.lower()
                if file_ext in ['.pdf', '.docx', '.doc', '.txt']:
                    docs = self.load_document(file_path, metadata)
                    all_documents.extend(docs)
        
        print(f"📚 디렉토리 로드 완료: {len(all_documents)}개 청크")
        return all_documents


if __name__ == "__main__":
    # 테스트
    loader = DocumentLoader(chunk_size=500, chunk_overlap=50)
    
    # 샘플 텍스트 파일 생성
    sample_text = """
    바이오헬스 산업 개요
    
    바이오헬스 산업은 생명공학 기술을 활용하여 인간의 건강과 삶의 질을 향상시키는 산업입니다.
    주요 분야로는 신약 개발, 의료기기, 디지털 헬스케어 등이 있습니다.
    
    mRNA 백신 기술
    
    mRNA 백신은 메신저 RNA를 이용하여 우리 몸의 세포가 특정 단백질을 생성하도록 지시합니다.
    이 기술은 COVID-19 팬데믹 동안 빠르게 발전하였으며, 향후 암 치료 등에도 활용될 전망입니다.
    """
    
    os.makedirs("./test_docs", exist_ok=True)
    with open("./test_docs/sample.txt", "w", encoding="utf-8") as f:
        f.write(sample_text)
    
    docs = loader.load_document("./test_docs/sample.txt", {"subject": "바이오헬스 기초"})
    
    print("\n=== 로드된 문서 ===")
    for i, doc in enumerate(docs):
        print(f"\n📄 청크 {i+1}:")
        print(f"내용: {doc.page_content[:100]}...")
        print(f"메타데이터: {doc.metadata}")
